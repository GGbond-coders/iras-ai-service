"""
IRAS AI Service — 简历文件解析模块。

根据上传文件 UUID 定位本地存储的简历文件，按扩展名分发到
对应的解析器，提取纯文本内容供 LLM 诊断使用。

支持的格式：
- .txt / .text: 直接读取（UTF-8 编码，非法字节用 replace 容错）
- .pdf: PyPDF2 提取每页文本
- .docx: python-docx 提取段落文本
- .doc: 暂不支持（旧 Word 格式需额外依赖）

核心函数：
- parse_resume(file_id): 根据 file_id 查找并解析简历文件，返回纯文本

文件定位策略：
遍历 UPLOAD_DIR 目录，找到以 {file_id} 开头的文件（如 "abc123.pdf"）。
使用前缀匹配而非精确匹配，因为存储时文件名为 {uuid}{原始扩展名}。
"""

import os
import logging
from config import UPLOAD_DIR

logger = logging.getLogger(__name__)


def parse_resume(file_id: str) -> str:
    """
    解析简历文件，返回纯文本内容。

    流程：
    1. 遍历上传目录，找到以 file_id 为前缀的文件
    2. 按扩展名分发到对应解析器（txt → 直接读，pdf → PyPDF2，docx → python-docx）
    3. 返回解析后的纯文本字符串

    参数：
        file_id: 上传简历时返回的文件 UUID（不含扩展名）

    返回：
        解析后的简历纯文本内容

    异常：
        FileNotFoundError: 上传目录中找不到以 file_id 开头的文件
        RuntimeError: 不支持的格式或解析失败（含具体原因）
    """
    upload_dir = os.path.abspath(UPLOAD_DIR)
    logger.info(f"[ResumeParser] Looking for file_id='{file_id}' in {upload_dir}")

    # 遍历上传目录，前缀匹配 file_id
    for filename in os.listdir(upload_dir):
        if filename.startswith(file_id):
            filepath = os.path.join(upload_dir, filename)
            ext = os.path.splitext(filename)[1].lower()
            logger.info(f"[ResumeParser] Found file: '{filename}' (ext={ext})")

            # 纯文本：直接读取，非法字节用 U+FFFD 替换
            if ext in (".txt", ".text"):
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
                logger.info(f"[ResumeParser] TXT parsed: {len(text)} chars")
                return text

            # PDF：使用 PyPDF2 逐页提取文本
            elif ext == ".pdf":
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(filepath)
                    page_count = len(reader.pages)
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                    logger.info(f"[ResumeParser] PDF parsed: {page_count} pages, {len(text)} chars")
                    return text
                except Exception:
                    logger.error(f"[ResumeParser] PDF parsing failed: {filename}")
                    raise RuntimeError(f"PDF parsing failed: {filename}")

            # DOCX：使用 python-docx 逐段落提取文本
            elif ext in (".doc", ".docx"):
                try:
                    from docx import Document
                    doc = Document(filepath)
                    para_count = len(doc.paragraphs)
                    text = "\n".join(p.text for p in doc.paragraphs)
                    logger.info(f"[ResumeParser] DOCX parsed: {para_count} paragraphs, {len(text)} chars")
                    return text
                except Exception:
                    logger.error(f"[ResumeParser] DOCX parsing failed: {filename}")
                    raise RuntimeError(f"DOCX parsing failed: {filename}")

            # 不支持的格式
            else:
                logger.error(f"[ResumeParser] Unsupported file type: {ext}")
                raise RuntimeError(f"Unsupported file type: {ext}")

    # 遍历完所有文件仍未找到匹配的 file_id
    logger.error(f"[ResumeParser] File not found for file_id='{file_id}'")
    raise FileNotFoundError(f"Uploaded file not found: {file_id}")
